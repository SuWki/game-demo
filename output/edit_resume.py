from pathlib import Path

import docx
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

src = Path(r"E:\codex\unity-learning\output\resume_working.docx")
out = Path(r"E:\codex\unity-learning\output\王友隆-游戏执行策划简历_项目版.docx")
doc = docx.Document(str(src))
FONT = "微软雅黑"


def set_east_asia(run, font=FONT):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)


def clear_para(p):
    for child in list(p._p):
        p._p.remove(child)


def write_para(p, text, *, bold=False, size=10, align=None):
    clear_para(p)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    set_east_asia(run)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    return p


def insert_after(paragraph, text="", *, bold=False, size=10, align=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = docx.text.paragraph.Paragraph(new_p, paragraph._parent)
    p.style = paragraph.style
    write_para(p, text, bold=bold, size=size, align=align)
    return p


def delete_para(p):
    el = p._element
    el.getparent().remove(el)


def norm(text):
    return text.strip().replace("\t", "").replace(" ", "")


def find_idx(label):
    wanted = label.replace(" ", "")
    for i, p in enumerate(doc.paragraphs):
        if norm(p.text) == wanted:
            return i
    raise ValueError(f"not found {label}")


def delete_between(start_heading, end_heading):
    paras = doc.paragraphs
    start = find_idx(start_heading)
    end = find_idx(end_heading)
    for p in paras[start + 1 : end]:
        delete_para(p)
    return doc.paragraphs[start]


project_heading = delete_between("项目经历", "个人能力")
project_lines = [
    ("短局节点推进自动射击原型｜系统/战斗/数值设计与 AI 协作落地", True, 10.5, None),
    ("展示视频：待上传 B 站链接（本轮已录制素材，上传后替换）", False, 9.0, None),
    ("交付物：可运行 Web 原型 / 设计文档 / 数值审计 / QA 脚本 / 开发复盘", False, 9.0, None),
    ("· 设计“手动走位 + 自动攻击 + 五轮短局 + 节点选择 + 三选一强化”的核心循环，明确普通战、精英战、生存战、Boss 战的关卡职责。", False, 9.4, None),
    ("· 设计暴击、穿透、穿梭三条路线，并按 starter / bridge / payoff / finisher 分层，控制路线牌不能一张成型。", False, 9.4, None),
    ("· 建立路线牌价值审计与实机反馈闭环，围绕普通强化价值、路线强度、Boss 安全区、敌人提示和表现层读局持续迭代。", False, 9.4, None),
    ("· 输出接手入口、设计基线、测试手册、开发复盘和 AI 协作规则，保证项目可被他人无上下文接手继续开发。", False, 9.4, None),
    ("灵兽系统改版方案（个人练习）｜系统改版与配表落地练习", True, 10.0, None),
    ("链接：https://b0l1lf6nijp.feishu.cn/wiki/BD3Nwf3YniY7y9kjoVNcPCnunOb", False, 9.0, None),
    ("交付物：规则文档 / 核心配表 / 测试用例", False, 9.0, None),
    ("· 基于《寻道大千》灵兽系统真实链路，梳理获取、养成、塑魂、内丹、上阵、回收等流程，补齐关键边界与异常场景。", False, 9.4, None),
    ("· 设计 9 张核心配置表，定义字段、ID 规则、关联关系与校验规则，并输出 4 类验收口径与 6 条测试用例示例。", False, 9.4, None),
    ("寻道大千系统分析与拆解项目｜系统循环与资源流向分析", True, 10.0, None),
    ("链接：https://ai.feishu.cn/wiki/FMDSwK3qti1BrCkwi1EcI9dGnyd", False, 9.0, None),
    ("交付物：系统流程图 / 核心玩法循环 / 资源流向分析", False, 9.0, None),
    ("· 拆解“砍树—养成—战斗”循环中的目标推进、反馈节点与资源消耗关系，整理灵兽、装备等核心系统的产出/消耗闭环。", False, 9.4, None),
    ("· 结合付费体验归纳首充、月卡、终身特权等付费点的触发时机、收益结构与心理驱动，并提出 PVP 公平性与后期消耗风险。", False, 9.4, None),
]
for text, bold, size, align in reversed(project_lines):
    insert_after(project_heading, text, bold=bold, size=size, align=align)

game_heading = delete_between("游戏经历", "工作经历")
game_lines = [
    ("定位：以策划视角拆解系统、战斗节奏与付费/留存逻辑，作为项目设计参考，而非单纯玩家经历。", False, 9.4, None),
    ("· 《APEX Legends》：长期关注英雄定位、武器数值、地图轮转与安全区节奏，拆解“搜刮—交战—转移—决赛圈”的压力曲线。", False, 9.4, None),
    ("· 《寻道大千》：拆解灵兽、精怪、装备、付费点与资源流向，整理养成目标、卡点、消耗与付费触发逻辑。", False, 9.4, None),
    ("· 输出方式：将体验结论转化为流程图、规则文档、配表字段和验收点，服务后续系统设计练习。", False, 9.4, None),
]
for text, bold, size, align in reversed(game_lines):
    insert_after(game_heading, text, bold=bold, size=size, align=align)

for p in doc.paragraphs:
    if p.text.strip().startswith("持续进行游戏系统拆解与执行策划练习"):
        write_para(
            p,
            "持续进行游戏系统拆解、短局玩法原型设计与执行策划练习，已输出规则文档、配置表、流程图、测试用例与可运行原型。",
            size=10,
        )

work_heading = doc.paragraphs[find_idx("工作经历")]
page_break_p = OxmlElement("w:p")
work_heading._p.addprevious(page_break_p)
break_para = docx.text.paragraph.Paragraph(page_break_p, work_heading._parent)
break_run = break_para.add_run()
break_run.add_break(WD_BREAK.PAGE)

for p in doc.paragraphs:
    txt = p.text.strip()
    for r in p.runs:
        set_east_asia(r)
    if txt.startswith("·"):
        for r in p.runs:
            r.font.size = Pt(9.4)
    if txt.startswith("链接：") or txt.startswith("交付物：") or txt.startswith("展示视频："):
        for r in p.runs:
            r.font.size = Pt(9.0)

doc.save(str(out))
print(str(out))
